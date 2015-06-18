from xlrd import open_workbook
from docx import Document
from docx.shared import Inches
import unidecode 
import glob, os
import re

execfile('global.py')

# this file has a function that processes one file, and returns a hash with all the values
def readFile(filename):
	book = open_workbook(filename,on_demand=True)
	for name in book.sheet_names():
    		if name.endswith('Lab'):
			sheet = book.sheet_by_name(name)
	d = dict()
	s = []
	drugs = [sheet.cell_value(3, 0),sheet.cell_value(3,1),sheet.cell_value(3,2)]
	for i in drugs:
		if i != "":
			s.append(i)	
	d["DRUG"] = "+".join(s)
	d["PT_INIT"] = sheet.cell_value(3, 5)
	d["PT_MRN"] = sheet.cell_value(4, 5)
	d["DATE"] = sheet.cell_value(5, 5)
	d["LAPTT_R"] = sheet.cell_value(11, 1)
	d["PTTMX_R"] = sheet.cell_value(14, 1)
	d["LTT_R"] = sheet.cell_value(12, 1)
	d["LTTHEP_R"] = sheet.cell_value(13, 1)
	d["LTTHEP_P"] = sheet.cell_value(13, 4)
	d["LAPTT_U"] = sheet.cell_value(11, 3)
	d["LTT_L"] = sheet.cell_value(12, 2)
	d["LTT_U"] = sheet.cell_value(12, 3)
	d["LTTHEP_U"] = sheet.cell_value(13, 3)
	d["PTTMX_U"] = sheet.cell_value(14, 3)
	d["LAPTT_P"] = sheet.cell_value(11, 4)
	d["LTT_P"] = sheet.cell_value(12, 4)
	d["LTTHEP_P"] = sheet.cell_value(13, 4)
	d["PTTMX_P"] = sheet.cell_value(14, 4)
	d["PNP_R"] = sheet.cell_value(20, 1)
	d["PNP_R_LYS"] = sheet.cell_value(20, 2)
	d["PNP_SD"] = sheet.cell_value(20, 4)	
	d["PNP_U"] = sheet.cell_value(20, 6)

	# have to fix this in excel sheet, or else PNP_SD defaults to 42.0
	if d["PNP_R"] == "":
		d["PNP_SD"] = 0
		d["PNP_R"]  = 0
	
	d["DRVVS_R"] = sheet.cell_value(25, 1)
	d["DRVVMX_R"] = sheet.cell_value(26, 1)
	d["DRVVC_R"] = sheet.cell_value(27, 1)
	d["PCTCO_R"] = sheet.cell_value(28, 1)
	d["DRVVS_U"] = sheet.cell_value(25, 3)
	d["DRVVMX_U"] = sheet.cell_value(26, 3)
	d["PCTCO_U"] = sheet.cell_value(28, 3)
	d["DRVVS_P"] = sheet.cell_value(25, 4)
	d["DRVVMX_P"] = sheet.cell_value(26, 4)
	d["PCTCO_P"] = sheet.cell_value(28, 4)
	d["DPTS_R"] = sheet.cell_value(33, 1)
	d["DPTMX_R"] = sheet.cell_value(34, 1)
	d["DPTC_R"] = sheet.cell_value(35, 1)
	d["DPTCOR_R"] = sheet.cell_value(36, 1)
	d["DPTS_U"] = sheet.cell_value(33, 3)
	d["DPTMX_U"] = sheet.cell_value(34, 3)
	d["DPTCOR_U"] = sheet.cell_value(36, 3)
	d["DPTS_P"] = sheet.cell_value(33, 4)
	d["DPTMX_P"] = sheet.cell_value(34, 4)
	d["DPTCOR_P"] = sheet.cell_value(36, 4)
	return d

def writeFile(txt, filename):
	document = Document()
	document.add_paragraph(txt)	
	document.save(filename)


def extractConclusion(filename):
	if filename is None:
		return ""

	f = open(os.path.join(REPORTDIR,filename))
	document = Document(f)
	s = []
	flag = False
	for p in document.paragraphs:
		cleanedText = unidecode.unidecode(p.text)
		cleanedText = cleanedText.replace("\n"," ")
		cleanedText = cleanedText.replace("\t"," ")
		if (cleanedText.startswith("Conclusion") or cleanedText.startswith("Note") or flag) and (not (cleanedText.startswith("Jonathan") or (cleanedText.startswith("Angela")))):
			flag = True
			if len(cleanedText)>0:
				s.append(cleanedText)
	f.close()	
	return "\n".join(s)

def findFiles(filename):
	strippedFilename = os.path.basename(filename)
	a = strippedFilename.split(" ");
	s = []
	for file in os.listdir(REPORTDIR):
		if (file.endswith(".docx") or (file.endswith(".doc"))) and file.startswith(a[0]+" "+a[1]):
			s.append(file)
	s.sort()
	return s

def findLastFile(filename):
	s = findFiles(filename)
	if len(s)>0:
		return s[len(s)-1]
	else:
		return None


